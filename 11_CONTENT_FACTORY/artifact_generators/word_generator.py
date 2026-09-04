# 11_CONTENT_FACTORY/artifact_generators/word_generator.py — 真实 DOCX 生成

from __future__ import annotations

from pathlib import Path
from typing import Any


def generate_word(title: str, sections: list[dict], output_path: Path) -> dict[str, Any]:
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError as exc:
        return {"status": "error", "file_path": "", "error": f"python-docx not installed: {exc}"}

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    doc.add_heading(title, 0)

    for sec in sections:
        doc.add_heading(sec.get("title", "章节"), level=1)
        for para in sec.get("paragraphs", []):
            p = doc.add_paragraph(para)
            p.paragraph_format.space_after = Pt(6)
        if sec.get("table"):
            table_data = sec["table"]
            if table_data:
                tbl = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
                tbl.style = "Table Grid"
                for r, row in enumerate(table_data):
                    for c, val in enumerate(row):
                        tbl.rows[r].cells[c].text = str(val)

    doc.add_paragraph("")
    doc.add_paragraph("— AI Factory OS Content Factory 生成")
    doc.save(str(output_path))
    return {"status": "ok", "file_path": str(output_path)}
